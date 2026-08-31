from __future__ import annotations

from pathlib import Path
import zipfile


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DIR = ROOT / "sample_files"

SAMPLE_LINES = [
    "홍길동",
    "010-1234-5678",
    "teacher@school.kr",
    "900101-1234567",
    "1990.01.01",
    "서울특별시 강남구 역삼동",
    "123-456-789012",
    "니코초등학교",
]


def create_xlsx_sample() -> None:
    try:
        from openpyxl import Workbook
    except ImportError:
        print("skip xlsx: openpyxl is not installed")
        return

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sheet1"
    for row_index, value in enumerate(SAMPLE_LINES, start=1):
        worksheet.cell(row=row_index, column=1, value=value)
    worksheet["B1"] = "홍길동 / 010-1234-5678 / teacher@school.kr"
    worksheet["B2"] = "=CONCAT(A1,A2)"
    workbook.save(SAMPLE_DIR / "sample_pii.xlsx")

    empty = Workbook()
    empty.active.title = "Empty"
    empty.save(SAMPLE_DIR / "empty.xlsx")


def create_docx_sample() -> None:
    try:
        from docx import Document
    except ImportError:
        print("skip docx: python-docx is not installed")
        return

    document = Document()
    document.add_paragraph("홍길동 연락처는 010-1234-5678이고 이메일은 teacher@school.kr입니다.")
    table = document.add_table(rows=5, cols=1)
    table.cell(0, 0).text = "900101-1234567"
    table.cell(1, 0).text = "1990.01.01"
    table.cell(2, 0).text = "서울특별시 강남구 역삼동"
    table.cell(3, 0).text = "123-456-789012"
    table.cell(4, 0).text = "니코초등학교"
    document.save(SAMPLE_DIR / "sample_pii.docx")

    empty = Document()
    empty.add_paragraph("")
    empty.save(SAMPLE_DIR / "empty.docx")


def create_text_pdf_sample() -> None:
    try:
        import fitz
    except ImportError:
        print("skip text pdf: pymupdf is not installed")
        return

    document = fitz.open()
    page = document.new_page()
    for index, value in enumerate(SAMPLE_LINES):
        page.insert_text((72, 72 + index * 28), value)
    document.save(SAMPLE_DIR / "sample_pii_text.pdf")
    document.close()


def create_scan_like_pdf_sample() -> None:
    try:
        import fitz
        from PIL import Image, ImageDraw
    except ImportError:
        print("skip scan-like pdf: pymupdf or pillow is not installed")
        return

    image_path = SAMPLE_DIR / "scan_like_source.png"
    image = Image.new("RGB", (1000, 700), "white")
    draw = ImageDraw.Draw(image)
    for index, value in enumerate(SAMPLE_LINES):
        draw.text((60, 60 + index * 55), value, fill="black")
    image.save(image_path)

    document = fitz.open()
    page = document.new_page(width=1000, height=700)
    page.insert_image(page.rect, filename=str(image_path))
    document.save(SAMPLE_DIR / "sample_pii_scan_like.pdf")
    document.close()


def create_hwpx_sample() -> None:
    section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<root>
  <body>
    <p>홍길동 연락처는 010-1234-5678이고 이메일은 teacher@school.kr입니다.</p>
    <p>주민등록번호는 900101-1234567입니다.</p>
    <p>생년월일은 1990.01.01입니다.</p>
    <p>주소는 서울특별시 강남구 역삼동입니다.</p>
    <p>계좌번호는 123-456-789012입니다.</p>
    <p>학교명은 니코초등학교입니다.</p>
  </body>
</root>
"""
    with zipfile.ZipFile(SAMPLE_DIR / "sample_pii.hwpx", "w") as package:
        package.writestr("mimetype", "application/hwp+zip")
        package.writestr("Contents/section0.xml", section_xml)
        package.writestr("Contents/header.xml", "<root><meta>sample</meta></root>")


def create_error_samples() -> None:
    (SAMPLE_DIR / "unsupported.txt").write_text("\n".join(SAMPLE_LINES), encoding="utf-8")
    (SAMPLE_DIR / "corrupted.hwpx").write_bytes(b"not a valid zip package")
    (SAMPLE_DIR / "corrupted.pdf").write_bytes(b"not a valid pdf")


def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    create_xlsx_sample()
    create_docx_sample()
    create_text_pdf_sample()
    create_scan_like_pdf_sample()
    create_hwpx_sample()
    create_error_samples()
    print(f"sample files created in: {SAMPLE_DIR}")


if __name__ == "__main__":
    main()
