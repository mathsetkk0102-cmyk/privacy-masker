from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.file_utils import make_masked_output_path
from core.models import DetectionItem
from core.pii_detector import PiiDetector
from processors.base_processor import BaseProcessor


class DocxProcessor(BaseProcessor):
    file_type = "docx"
    supported_suffixes = {".docx"}

    def __init__(self, detector: PiiDetector | None = None) -> None:
        self.detector = detector or PiiDetector()

    def detect(self, file_path: str | Path) -> list[DetectionItem]:
        input_path = Path(file_path)
        document = Document(str(input_path))
        detections: list[DetectionItem] = []

        for index, paragraph in enumerate(document.paragraphs):
            detections.extend(self._detect_paragraph(input_path, paragraph, "본문", f"paragraph[{index}]"))

        for table_index, table in enumerate(document.tables):
            detections.extend(self._detect_table(input_path, table, f"table[{table_index}]"))

        return detections

    def apply_masking(self, file_path: str | Path, output_dir: str | Path, items: list[DetectionItem]) -> str:
        input_path = Path(file_path)
        output_folder = Path(output_dir)
        output_path = make_masked_output_path(input_path, output_folder)
        document = Document(str(input_path))

        self._apply_to_document(document, input_path, items)
        output_folder.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return str(output_path)

    def save_masked_copy(self, input_path: Path, output_path: Path, detections: list[DetectionItem]) -> None:
        document = Document(str(input_path))
        self._apply_to_document(document, input_path, detections)
        document.save(str(output_path))

    def _detect_paragraph(
        self,
        input_path: Path,
        paragraph: Paragraph,
        page_or_sheet: str,
        location: str,
    ) -> list[DetectionItem]:
        if not paragraph.text:
            return []
        source = {
            "file_path": str(input_path),
            "file_type": self.file_type,
            "page_or_sheet": page_or_sheet,
            "location": location,
        }
        return self.detector.detect_text(paragraph.text, source)

    def _detect_table(self, input_path: Path, table: Table, table_label: str) -> list[DetectionItem]:
        detections: list[DetectionItem] = []
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    location = f"{table_label}.row[{row_index}].cell[{cell_index}].paragraph[{paragraph_index}]"
                    detections.extend(self._detect_paragraph(input_path, paragraph, "표", location))
                for nested_index, nested_table in enumerate(cell.tables):
                    nested_label = f"{table_label}.row[{row_index}].cell[{cell_index}].table[{nested_index}]"
                    detections.extend(self._detect_table(input_path, nested_table, nested_label))
        return detections

    def _apply_to_document(self, document: DocumentObject, input_path: Path, items: list[DetectionItem]) -> None:
        grouped_items = self._group_selected_items(items, input_path)

        for index, paragraph in enumerate(document.paragraphs):
            key = ("본문", f"paragraph[{index}]")
            if key in grouped_items:
                # 첫 버전에서는 paragraph.text 치환을 사용한다. run 단위 서식은 일부 단순화될 수 있다.
                paragraph.text = self._replace_selected(paragraph.text, grouped_items[key])

        for table_index, table in enumerate(document.tables):
            self._apply_to_table(table, f"table[{table_index}]", grouped_items)

    def _apply_to_table(
        self,
        table: Table,
        table_label: str,
        grouped_items: dict[tuple[str, str], list[DetectionItem]],
    ) -> None:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    key = ("표", f"{table_label}.row[{row_index}].cell[{cell_index}].paragraph[{paragraph_index}]")
                    if key in grouped_items:
                        # 첫 버전에서는 paragraph.text 치환을 사용한다. run 단위 서식은 일부 단순화될 수 있다.
                        paragraph.text = self._replace_selected(paragraph.text, grouped_items[key])
                for nested_index, nested_table in enumerate(cell.tables):
                    nested_label = f"{table_label}.row[{row_index}].cell[{cell_index}].table[{nested_index}]"
                    self._apply_to_table(nested_table, nested_label, grouped_items)

    @staticmethod
    def _replace_selected(text: str, selected_items: list[DetectionItem]) -> str:
        result = text
        for item in selected_items:
            result = result.replace(item.original_text, item.masked_text, 1)
        return result

    @staticmethod
    def _group_selected_items(items: list[DetectionItem], input_path: Path) -> dict[tuple[str, str], list[DetectionItem]]:
        grouped: dict[tuple[str, str], list[DetectionItem]] = defaultdict(list)
        for item in items:
            if not item.selected:
                continue
            if Path(item.file_path) != input_path:
                continue
            grouped[(item.page_or_sheet, item.location)].append(item)
        return grouped
