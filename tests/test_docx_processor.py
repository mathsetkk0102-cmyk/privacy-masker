from pathlib import Path

from docx import Document

from processors.docx_processor import DocxProcessor


def create_sample_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("홍길동 학생의 연락처는 010-1234-5678이고 이메일은 teacher@school.kr입니다.")
    table = document.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "900101-1234567"
    table.cell(1, 0).text = "서울특별시 강남구 역삼동"
    table.cell(2, 0).text = "니코초등학교"
    document.save(path)


def read_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_docx_detects_pii_with_paragraph_and_table_locations(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    create_sample_docx(input_path)

    items = DocxProcessor().detect(input_path)
    types = {item.pii_type for item in items}

    assert {"이름", "전화번호", "이메일", "주민등록번호", "주소", "학교명"}.issubset(types)
    assert any(item.page_or_sheet == "본문" and item.location.startswith("paragraph[") for item in items)
    assert any(item.page_or_sheet == "표" and item.location.startswith("table[") for item in items)

    by_type = {item.pii_type: item for item in items}
    assert by_type["전화번호"].review_status == "자동선택"
    assert by_type["이메일"].review_status == "자동선택"
    assert by_type["이름"].review_status == "확인 필요"
    assert by_type["주소"].review_status == "확인 필요"
    assert by_type["학교명"].review_status == "확인 필요"


def test_docx_apply_masking_only_selected_items_and_keeps_original(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    output_dir = tmp_path / "out"
    create_sample_docx(input_path)

    processor = DocxProcessor()
    items = processor.detect(input_path)
    for item in items:
        item.selected = item.pii_type == "전화번호"

    output_path = Path(processor.apply_masking(input_path, output_dir, items))

    original_text = read_docx_text(input_path)
    result_text = read_docx_text(output_path)

    assert output_path.exists()
    assert output_path.name == "masked_sample.docx"
    assert "010-1234-5678" in original_text
    assert "010-****-5678" in result_text
    assert "teacher@school.kr" in result_text
    assert "니코초등학교" in result_text


def test_docx_uses_user_edited_masked_text(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    output_dir = tmp_path / "out"
    create_sample_docx(input_path)

    processor = DocxProcessor()
    items = processor.detect(input_path)
    for item in items:
        item.selected = item.original_text == "teacher@school.kr"
        if item.selected:
            item.masked_text = "직접수정@example.com"

    output_path = Path(processor.apply_masking(input_path, output_dir, items))
    result_text = read_docx_text(output_path)

    assert "직접수정@example.com" in result_text
